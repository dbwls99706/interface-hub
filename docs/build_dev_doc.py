"""docs/개발문서.docx 생성 스크립트."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parent
OUTPUT = DOCS_DIR / "개발문서.docx"

KOR_FONT = "맑은 고딕"
CODE_FONT = "Consolas"


def set_korean_font(run, size_pt: float, bold: bool = False, color=None) -> None:
    run.font.name = KOR_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), KOR_FONT)
    r_fonts.set(qn("w:ascii"), KOR_FONT)
    r_fonts.set(qn("w:hAnsi"), KOR_FONT)


def set_code_font(run, size_pt: float = 9) -> None:
    run.font.name = CODE_FONT
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), CODE_FONT)
    r_fonts.set(qn("w:ascii"), CODE_FONT)
    r_fonts.set(qn("w:hAnsi"), CODE_FONT)


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: float = 11,
    bold: bool = False,
    align=None,
    space_after: float = 6,
    color=None,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_korean_font(run, size, bold=bold, color=color)
    return p


def add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    if level == 0:
        size = 22
    elif level == 1:
        size = 16
    elif level == 2:
        size = 13
    else:
        size = 12
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_korean_font(run, size, bold=True, color=RGBColor(0x1F, 0x2D, 0x3D))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.35
        if p.runs:
            run = p.runs[0]
            run.text = item
        else:
            run = p.add_run(item)
        set_korean_font(run, 11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.35
        if p.runs:
            run = p.runs[0]
            run.text = item
        else:
            run = p.add_run(item)
        set_korean_font(run, 11)


def add_code_block(doc: Document, code: str) -> None:
    """회색 배경 + Consolas 9pt 코드 블록을 1x1 표로 구현."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(15.8)
    cell = table.cell(0, 0)
    cell.width = Cm(15.8)
    _set_cell_shading(cell, "F4F5F7")
    cell.text = ""
    lines = code.splitlines() or [""]
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(line if line else " ")
        set_code_font(run, 9)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    col_widths_cm: list[float] | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    if col_widths_cm:
        for col_idx, width in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[col_idx].width = Cm(width)

    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_korean_font(run, 10.5, bold=True)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            set_korean_font(run, 10.5)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def build_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    add_paragraph(
        doc,
        "Interface Hub 개발문서",
        size=26,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x10, 0x2A, 0x43),
        space_after=4,
    )
    add_paragraph(
        doc,
        "아키텍처와 바이브코딩 프로세스",
        size=15,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=24,
    )
    add_paragraph(
        doc,
        "노아에이티에스 2025년 상반기 연구소 인력 채용 포트폴리오",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x44, 0x55, 0x66),
        space_after=80,
    )
    add_paragraph(
        doc,
        "작성자: 홍유진 (GitHub: dbwls99706)",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "작성일: 2026-04-20",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "라이브 데모: https://interface-hub.vercel.app",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "저장소: https://github.com/dbwls99706/interface-hub",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_page_break()


def build_toc(doc: Document) -> None:
    add_heading(doc, "목차", level=0)
    items = [
        "1. 문서 목적과 범위",
        "2. 시스템 아키텍처",
        "3. 데이터 모델",
        "4. Adapter 패턴 설계",
        "5. 실행 엔진",
        "6. 프런트엔드 데이터 흐름",
        "7. 로컬 SQLite와 Turso 전환 전략",
        "8. 바이브코딩 프로세스",
        "9. 검증 전략",
        "10. 트러블슈팅 사례",
        "11. 부록",
    ]
    for item in items:
        add_paragraph(doc, item, size=12, space_after=4)
    doc.add_page_break()


def section_scope(doc: Document) -> None:
    add_heading(doc, "1. 문서 목적과 범위", level=1)
    add_paragraph(
        doc,
        "이 문서는 Interface Hub의 설계 결정과 개발 프로세스를 심사위원이 재현 가능한 수준으로 기술한다. "
        "제품의 기능 명세는 기획서(docs/기획서.docx)를 참조한다.",
    )
    add_bullets(
        doc,
        [
            "대상 독자: 채용 심사위원, 신규 합류 개발자.",
            "다루는 범위: 시스템 아키텍처, 데이터 모델, 주요 패턴, 개발 프로세스, 트러블슈팅.",
            "다루지 않는 범위: 사용자 가이드, UI 카피 리뷰, 화면별 인터랙션 세부 스펙.",
        ],
    )


def section_architecture(doc: Document) -> None:
    add_heading(doc, "2. 시스템 아키텍처", level=1)

    add_heading(doc, "2.1 고수준 구성", level=2)
    add_paragraph(
        doc,
        "브라우저는 Next.js 서버와 Server Action으로 통신한다. 서버는 Prisma Client를 통해 libSQL driver adapter로 DB에 접근한다. "
        "로컬 개발에서는 SQLite 파일, 프로덕션에서는 Turso(libSQL) 인스턴스를 사용한다.",
    )
    add_code_block(
        doc,
        "Browser\n"
        "   │  fetch / form submit\n"
        "   ▼\n"
        "Next.js (Server Component, Server Action)\n"
        "   │  prisma.* 호출\n"
        "   ▼\n"
        "Prisma Client (PrismaLibSql adapter)\n"
        "   │\n"
        "   ├─ 로컬: file:./dev.db (SQLite)\n"
        "   └─ 프로덕션: libsql://...turso.io (Turso)",
    )

    add_heading(doc, "2.2 주요 결정 사항", level=2)
    add_table(
        doc,
        headers=["결정 지점", "선택", "이유"],
        rows=[
            [
                "App Router vs Pages Router",
                "App Router",
                "Server Component와 Server Action으로 API Route를 제거하고 타입 경로를 단일화한다.",
            ],
            [
                "Server Action vs API Route",
                "Server Action 우선",
                "클라이언트와 서버 사이 타입 공유와 입력 검증 경로를 하나로 묶는다.",
            ],
            [
                "REST + Mock 혼합",
                "MVP 전략",
                "REST 실동작으로 실증하고 나머지 프로토콜은 Mock으로 확장 구조를 증명한다.",
            ],
            [
                "Prisma + libSQL Adapter",
                "Turso 전환 용이",
                "런타임 환경변수 하나로 로컬 SQLite와 Turso를 분기한다.",
            ],
            [
                "SWR 폴링 vs WebSocket",
                "SWR 조건부 폴링",
                "서버리스 친화적이고 RUNNING일 때만 폴링을 활성화해 비용을 낮춘다.",
            ],
        ],
        col_widths_cm=[4.2, 3.8, 7.8],
    )


def section_data_model(doc: Document) -> None:
    add_heading(doc, "3. 데이터 모델", level=1)

    add_heading(doc, "3.1 엔티티", level=2)
    add_paragraph(doc, "3개 엔티티로 전 프로토콜의 운영 데이터를 담는다.")
    add_table(
        doc,
        headers=["엔티티", "주요 필드"],
        rows=[
            [
                "Interface",
                "id, name, protocol, endpoint, config, description, isActive, createdAt, updatedAt",
            ],
            [
                "Execution",
                "id, interfaceId, status, startedAt, finishedAt, durationMs, errorMessage, request, response, retryOfId",
            ],
            [
                "ExecutionLog",
                "id, executionId, level, message, metadata, loggedAt",
            ],
        ],
        col_widths_cm=[3.0, 12.8],
    )

    add_heading(doc, "3.2 관계", level=2)
    add_bullets(
        doc,
        [
            "Interface 1:N Execution (인터페이스 삭제 시 Execution Cascade).",
            "Execution 1:N ExecutionLog (실행 삭제 시 로그 Cascade).",
            "Execution 1:N Execution (retryOfId로 재처리 체인을 구성한다).",
        ],
    )

    add_heading(doc, "3.3 주요 인덱스", level=2)
    add_table(
        doc,
        headers=["엔티티", "인덱스", "이유"],
        rows=[
            ["Execution", "(interfaceId, startedAt)", "인터페이스별 최신 실행 조회"],
            ["Execution", "(status)", "상태 필터(실패, 진행 중)"],
            ["ExecutionLog", "(executionId, loggedAt)", "로그 타임라인 정렬"],
        ],
        col_widths_cm=[3.2, 5.6, 7.0],
    )

    add_heading(doc, "3.4 SQLite 제약 대응", level=2)
    add_bullets(
        doc,
        [
            "JSON 필드는 TEXT로 저장하고 parseJson/stringifyJson 헬퍼로 직렬화한다.",
            "enum은 사용하지 않고 literal union 타입과 Zod refine으로 런타임 검증한다.",
            "Turso의 libSQL은 SQLite 상위 호환이므로 동일 스키마가 그대로 동작한다.",
        ],
    )


def section_adapter(doc: Document) -> None:
    add_heading(doc, "4. Adapter 패턴 설계", level=1)

    add_heading(doc, "4.1 인터페이스 정의", level=2)
    add_code_block(
        doc,
        "export type InterfaceAdapter = {\n"
        "  protocol: Protocol\n"
        "  validateConfig(config: unknown):\n"
        "    | { ok: true; config: Record<string, unknown> }\n"
        "    | { ok: false; error: string }\n"
        "  execute(ctx: AdapterContext): Promise<AdapterResult>\n"
        "}",
    )
    add_paragraph(
        doc,
        "모든 어댑터는 동일 시그니처를 구현한다. 실행 엔진은 Protocol 문자열로 어댑터를 조회해 호출할 뿐, "
        "프로토콜별 분기를 갖지 않는다.",
    )

    add_heading(doc, "4.2 구현 구조", level=2)
    add_bullets(
        doc,
        [
            "lib/adapters/types.ts: AdapterContext, AdapterResult, AdapterLog, InterfaceAdapter 공용 타입.",
            "lib/adapters/_logger.ts: 어댑터 내부에서 로그를 수집하는 createLogger().",
            "lib/adapters/rest.ts: 실제 fetch 구현. Zod로 config 검증, AbortController와 signal 병합, content-type 분기.",
            "lib/adapters/mock.ts: Mock 베이스인 createMockAdapter(). step 진행률에 맞춰 로그를 쌓고 signal.aborted를 50ms 주기로 점검.",
            "lib/adapters/{soap,mq,batch,sftp}.ts: 프로토콜별 Mock 어댑터. 고유한 step 시퀀스, 지연 범위, 에러 메시지, 샘플 request/response 정의.",
            "lib/adapters/registry.ts: Protocol → InterfaceAdapter 맵과 getAdapter() 함수.",
        ],
    )

    add_heading(doc, "4.3 확장성 검증", level=2)
    add_paragraph(
        doc,
        "신규 프로토콜 추가는 어댑터 파일 1개와 registry.ts의 1줄로 끝나도록 설계했다. 예를 들어 gRPC 어댑터를 추가한다면 다음 순서로 진행한다.",
    )
    add_numbered(
        doc,
        [
            "lib/adapters/grpc.ts 생성 후 InterfaceAdapter를 구현한다. 내부는 createMockAdapter를 쓰거나 실 gRPC client를 감싸면 된다.",
            "lib/types/db.ts의 Protocol 상수에 'GRPC'를 추가한다.",
            "lib/adapters/registry.ts의 맵에 GRPC: grpcAdapter 한 줄을 추가한다.",
            "UI의 프로토콜 Select와 ProtocolBadge 스타일에 GRPC 하나만 보강하면 전체 파이프라인(실행 엔진, 실행 이력, 대시보드)이 자동으로 해당 프로토콜을 인지한다.",
        ],
    )

    add_heading(doc, "4.4 오류 수렴 원칙", level=2)
    add_bullets(
        doc,
        [
            "모든 어댑터는 throw하지 않는다.",
            "모든 에러는 AdapterResult.status = 'FAILED'와 errorMessage로 수렴한다.",
            "이유: 실행 엔진의 never-throw 정책을 보장하고, 부분 실패가 Execution 레코드 누락으로 이어지지 않도록 한다.",
        ],
    )


def section_engine(doc: Document) -> None:
    add_heading(doc, "5. 실행 엔진", level=1)

    add_heading(doc, "5.1 실행 시퀀스", level=2)
    add_numbered(
        doc,
        [
            "executeInterface(interfaceId)를 호출한다 (Server Action).",
            "prisma.interface.findUnique로 Interface를 조회하고 isActive를 확인한다.",
            "prisma.execution.create로 Execution을 PENDING 상태로 생성한다.",
            "runExecution()에서 status를 RUNNING으로 업데이트하고 startedAt을 갱신한다.",
            "AbortController를 만들고 setTimeout으로 30초 타임아웃을 건다.",
            "getAdapter(protocol)으로 어댑터를 가져와 validateConfig → execute 순으로 호출한다.",
            "Execution을 최종 status와 durationMs, request, response로 업데이트한다.",
            "ExecutionLog를 prisma.executionLog.createMany로 일괄 삽입한다.",
            "revalidatePath로 /interfaces/:id와 /executions의 캐시를 무효화한다.",
        ],
    )

    add_heading(doc, "5.2 Retry 처리", level=2)
    add_bullets(
        doc,
        [
            "원본 Execution은 상태를 유지한다.",
            "새 Execution을 생성하고 retryOfId로 원본 Execution의 id를 참조한다.",
            "이후 runExecution()으로 실행 흐름을 위임해 코드를 재사용한다.",
        ],
    )

    add_heading(doc, "5.3 타임아웃과 취소", level=2)
    add_bullets(
        doc,
        [
            "AbortSignal이 fetch와 Mock 어댑터의 sleep 양쪽으로 전파된다.",
            "Mock의 sleep은 50ms 주기로 signal.aborted를 체크한 뒤 즉시 종료한다.",
            "취소된 경우에도 Execution은 FAILED로 마무리되어 recon 지점이 유지된다.",
        ],
    )


def section_frontend(doc: Document) -> None:
    add_heading(doc, "6. 프런트엔드 데이터 흐름", level=1)

    add_heading(doc, "6.1 화면별 렌더링 전략", level=2)
    add_table(
        doc,
        headers=["화면", "전략", "이유"],
        rows=[
            ["/interfaces", "Server Component", "정적 목록. revalidatePath로 캐시 갱신 충분."],
            ["/interfaces/new, /edit", "Server Page + Client Form", "RHF 인터랙션은 Client로 분리."],
            [
                "/executions",
                "Server Component + Client 리스트",
                "필터는 URL params, 데이터는 useSWRInfinite.",
            ],
            [
                "/executions/[id]",
                "Server 초기 로드 + Client SWR",
                "RUNNING 상태에서 3초 폴링으로 실시간 반영.",
            ],
            [
                "/dashboard",
                "Server Component",
                "집계는 SSR에서 수행, 차트만 Client.",
            ],
        ],
        col_widths_cm=[3.8, 4.6, 7.4],
    )

    add_heading(doc, "6.2 조건부 폴링 (SWR)", level=2)
    add_paragraph(
        doc,
        "useSWR의 refreshInterval을 함수형으로 제공하면 최신 데이터 내용을 보고 간격을 동적으로 결정할 수 있다. "
        "RUNNING이나 PENDING이 포함된 경우에만 3초로 폴링하고, 그 외에는 0(비활성)으로 내려 서버리스 호출 비용을 줄인다.",
    )
    add_code_block(
        doc,
        "useSWRInfinite<ListExecutionsResult>(\n"
        "  getKey,\n"
        "  ([, , , , cursor]) => listExecutions({ ...filters, cursor }),\n"
        "  {\n"
        "    refreshInterval: (latest) =>\n"
        "      latest?.some((p) =>\n"
        "        p.items.some((x) => isLiveStatus(x.status)),\n"
        "      )\n"
        "        ? 3000\n"
        "        : 0,\n"
        "    revalidateOnFocus: false,\n"
        "    revalidateFirstPage: true,\n"
        "  },\n"
        ");",
    )

    add_heading(doc, "6.3 useSWRInfinite 더 보기", level=2)
    add_bullets(
        doc,
        [
            "getKey에 filters와 커서를 포함시켜 필터 변경 시 자동으로 캐시를 분리한다.",
            "revalidateFirstPage: true로 첫 페이지는 항상 최신 상태를 유지한다.",
            "더 보기 버튼은 lastPage.nextCursor가 있을 때만 노출한다.",
        ],
    )


def section_turso(doc: Document) -> None:
    add_heading(doc, "7. 로컬 SQLite와 Turso 전환 전략", level=1)

    add_heading(doc, "7.1 문제", level=2)
    add_bullets(
        doc,
        [
            "Prisma 7 CLI는 libsql:// URL을 직접 이해하지 못한다.",
            "Driver adapter는 런타임에만 주입되고, migrate 커맨드에 공식 경로로 전달할 수 없다.",
            "Vercel 배포 시 Turso에 마이그레이션을 어떻게 적용할지가 별도 과제가 된다.",
        ],
    )

    add_heading(doc, "7.2 해결", level=2)
    add_bullets(
        doc,
        [
            "런타임은 lib/prisma.ts에서 PrismaLibSql이 file:과 libsql: 양쪽을 처리한다.",
            "환경변수 우선순위는 TURSO_DATABASE_URL > DATABASE_URL(libsql://) > DATABASE_URL(file:) 순이다.",
            "마이그레이션은 scripts/turso-deploy.ts를 통해 우회한다. libSQL client로 SQL을 직접 실행한다.",
        ],
    )

    add_heading(doc, "7.3 Turso 배포 스크립트 원리", level=2)
    add_bullets(
        doc,
        [
            "_applied_migrations 테이블을 만들어 멱등성을 확보한다.",
            "prisma/migrations를 이름 순으로 순회하고, 이미 적용된 마이그레이션은 스킵한다.",
            "주석 라인을 라인 단위로 제거한 뒤 세미콜론 기준으로 statement를 분리해 하나씩 execute한다.",
            "실패 시 롤백 대신 scripts/turso-reset.ts로 적용 기록을 초기화하고 재시도한다.",
        ],
    )
    add_code_block(
        doc,
        "const stripComments = (sql: string): string =>\n"
        "  sql\n"
        "    .split(\"\\n\")\n"
        "    .filter((line) => !line.trim().startsWith(\"--\"))\n"
        "    .join(\"\\n\");\n"
        "\n"
        "const splitStatements = (sql: string): string[] =>\n"
        "  stripComments(sql)\n"
        "    .split(/;\\s*$/m)\n"
        "    .map((s) => s.trim())\n"
        "    .filter((s) => s.length > 0);",
    )


def section_vibe(doc: Document) -> None:
    add_heading(doc, "8. 바이브코딩 프로세스", level=1)

    add_heading(doc, "8.1 접근 방식", level=2)
    add_bullets(
        doc,
        [
            "Claude Code를 페어 프로그래밍 파트너로 두고 1일 MVP를 목표로 설정했다.",
            "작업을 Phase 0부터 Phase 8까지 쪼개고, 각 Phase의 산출물과 검증 기준을 먼저 문서화했다.",
            "구현은 프롬프트 엔지니어링, 코드 리뷰, 검증 루프를 빠르게 반복하는 방식으로 진행했다.",
        ],
    )

    add_heading(doc, "8.2 프롬프트 엔지니어링 원칙", level=2)

    add_paragraph(doc, "원칙 1: 산출물 명세를 타입 시그니처로 못박는다.", bold=True, space_after=2)
    add_paragraph(
        doc,
        "Server Action과 Adapter 인터페이스의 반환 타입을 프롬프트 최상단에 선언한 뒤 구현을 요청한다. "
        "AI가 파라미터 순서와 에러 형태를 추측하지 않고 명시된 계약을 정확히 따르게 된다.",
    )

    add_paragraph(doc, "원칙 2: 환경적 제약을 프롬프트 상단에 명시한다.", bold=True, space_after=2)
    add_paragraph(
        doc,
        "Next.js 16의 params Promise 타입, Prisma 7의 import 경로(@/lib/generated/prisma/client), "
        "사용 가능한 외부 패키지 목록을 프롬프트 서두에 둔다. 구버전 API로 코드를 생성하는 실수가 크게 줄어든다.",
    )

    add_paragraph(doc, "원칙 3: 하지 말아야 할 것까지 적는다.", bold=True, space_after=2)
    add_paragraph(
        doc,
        "useState 대신 useTransition, initialData는 이 Phase에서 제외, curl/bash 명령은 문서에만 등 "
        "금지 조건을 명시한다. AI가 의도치 않은 유연성을 발휘해 scope를 넘는 것을 막는다.",
    )

    add_paragraph(doc, "원칙 4: 검증 기준을 프롬프트에 내장한다.", bold=True, space_after=2)
    add_paragraph(
        doc,
        "예를 들어 'npx tsc --noEmit 에러 0 확인', '/executions에서 데이터가 보이는지 확인'처럼 "
        "완료 조건을 프롬프트에 포함시킨다. AI가 스스로 자체 점검 후 보고하게 된다.",
    )

    add_heading(doc, "8.3 실제 프롬프트 예시", level=2)

    add_paragraph(doc, "예시 1: Phase 3 Adapter 레이어 프롬프트 발췌", bold=True, space_after=2)
    add_code_block(
        doc,
        "lib/adapters/types.ts\n"
        "  export type InterfaceAdapter = {\n"
        "    protocol: Protocol\n"
        "    validateConfig(config: unknown):\n"
        "      | { ok: true; config: Record<string, unknown> }\n"
        "      | { ok: false; error: string }\n"
        "    execute(ctx: AdapterContext): Promise<AdapterResult>\n"
        "  }\n"
        "\n"
        "품질 기준:\n"
        "- 어댑터는 throw하지 않음. 모든 에러는 FAILED + errorMessage로 수렴\n"
        "- registry에 새 프로토콜 추가가 '어댑터 파일 1개 + 한 줄'로 끝나야 함",
    )
    add_paragraph(
        doc,
        "해설: 타입 시그니처를 먼저 선언하고 오류 수렴 원칙을 명문화했다. 이후 생성된 어댑터들은 시그니처 일관성과 "
        "never-throw 정책을 자동으로 따랐으며, 실행 엔진 쪽에서 try/catch 복잡도를 크게 줄일 수 있었다.",
    )

    add_paragraph(doc, "예시 2: Phase 4 실시간 폴링 요구 부분", bold=True, space_after=2)
    add_code_block(
        doc,
        "RUNNING이 포함된 경우만 3초 폴링, 아니면 비활성 (폴링 조건부)\n"
        "  * refreshInterval 함수형:\n"
        "    (data) => data?.items.some(x => x.status === 'RUNNING' || x.status === 'PENDING')\n"
        "      ? 3000 : 0",
    )
    add_paragraph(
        doc,
        "해설: '조건부 폴링'이라는 요구를 말이 아닌 함수 시그니처 수준으로 구체화했다. "
        "서버리스 환경에서 불필요한 호출을 막는 의도를 AI가 놓치지 않도록 한 것이다.",
    )

    add_paragraph(doc, "예시 3: Turso 전환 타이밍 이슈 해결 프롬프트", bold=True, space_after=2)
    add_code_block(
        doc,
        "원인: Prisma migration SQL이 '-- CreateTable\\nCREATE TABLE ...' 형태라\n"
        "statement 단위로 split한 뒤 startsWith('--') 필터가 CREATE 문을 포함한 덩어리를 통째로 제거함.\n"
        "\n"
        "수정 방향:\n"
        "1. SQL을 라인 단위로 먼저 순회하면서 '--'로 시작하는 라인은 제거\n"
        "2. 주석 제거된 SQL 덩어리를 세미콜론 기준으로 split",
    )
    add_paragraph(
        doc,
        "해설: 현상만 알려주지 않고 원인과 수정 방향을 함께 제시했다. AI가 버그 원인을 추측하는 데 드는 시간을 "
        "아끼고, 동일한 패턴 실수를 반복하지 않도록 원인을 명문화해 남겼다.",
    )

    add_heading(doc, "8.4 검증 루프 (1 Phase당 반복)", level=2)
    add_numbered(
        doc,
        [
            "요구 정의: 데이터 모델, 화면, 동작을 문장으로 정리한다.",
            "프롬프트: 타입 시그니처, 파일 경로, 제약 사항을 명시한다.",
            "생성: Claude Code가 파일을 생성하거나 수정한다.",
            "컴파일 검증: npx tsc --noEmit이 0이어야 한다.",
            "실동작 검증: 브라우저에서 실제 사용자 플로우를 수행한다.",
            "커밋: Conventional Commits 규칙으로 단일 논리 단위의 커밋 하나를 만든다.",
        ],
    )

    add_heading(doc, "8.5 AI가 틀렸을 때 대응", level=2)
    add_bullets(
        doc,
        [
            "사례 1: @prisma/adapter-libsql의 export 이름을 PrismaLibSQL로 잘못 주장 → d.ts 파일을 grep으로 직접 확인해 실제 export가 PrismaLibSql임을 근거로 지정.",
            "사례 2: SQL 주석 라인 필터가 CREATE 문을 통째로 제거하는 버그 → migration.sql 원문을 직접 열어 파싱 단계를 역추적 후 수정 방향을 프롬프트에 제시.",
            "대응 원칙: AI의 자기 확신에 끌려가지 않고 원천 소스(.d.ts, 실제 출력)를 확인한다.",
        ],
    )

    add_heading(doc, "8.6 소요 시간과 산출", level=2)
    add_table(
        doc,
        headers=["항목", "수치"],
        rows=[
            ["총 소요 시간", "약 8시간 (Phase 0 → Phase 8)"],
            ["생성/수정 파일", "약 40개 (TypeScript, Prisma, 스크립트 포함)"],
            ["커밋 수", "10~15개 (Phase 단위 + 버그 수정)"],
            ["프롬프트 → 수정 반복", "Phase당 평균 1~2회"],
        ],
        col_widths_cm=[5.0, 10.8],
    )

    add_heading(doc, "8.7 레슨의 전파: CLAUDE.md를 통한 반복 실수 제거", level=2)
    add_bullets(
        doc,
        [
            "Interface Hub 개발 중 Prisma 7, Recharts, shadcn/ui, Turso 등 여러 지점에서 예상치 못한 이슈를 만났다.",
            "단순히 해결하고 넘어가는 대신, 각 이슈와 해결 방식을 CLAUDE.md에 '프로젝트별 환경 메모' 형식으로 기록했다.",
            "이후 같은 기술 스택으로 자매 프로젝트 Cost Compass를 시작할 때 CLAUDE.md를 복사하고 프로젝트 컨텍스트만 교체했다.",
            "결과: Interface Hub에서 겪은 11개 이슈 중 9개가 Cost Compass Phase 0~3에서 재발하지 않았다. Phase 1 DB 셋업은 Interface Hub보다 약 절반 시간에 완료됐다.",
            "이것이 바이브코딩이 단순한 AI 코드 생성과 다른 이유다. 프롬프트 엔지니어링만큼이나 CLAUDE.md 같은 컨텍스트 문서의 설계가 중요하다.",
        ],
    )
    add_paragraph(
        doc,
        "AI는 기억하지 못하지만, 개발자가 남긴 CLAUDE.md는 다음 세션에서 다시 읽힌다.",
        bold=True,
    )


def section_validation(doc: Document) -> None:
    add_heading(doc, "9. 검증 전략", level=1)

    add_heading(doc, "9.1 타입 안전", level=2)
    add_bullets(
        doc,
        [
            "TypeScript strict 모드를 전 코드에 적용한다.",
            "any 사용을 금지하고 unknown을 받은 뒤 타입 가드로 좁힌다.",
            "Zod 스키마는 react-hook-form 검증과 Server Action 입력 검증 양쪽에서 재사용한다.",
        ],
    )

    add_heading(doc, "9.2 런타임 검증", level=2)
    add_bullets(
        doc,
        [
            "Server Action은 never-throw 계약을 갖는다. 결과는 ActionResult<T> 판별 유니온으로 수렴한다.",
            "Adapter는 FAILED + errorMessage로 에러를 수렴해 실행 엔진이 Execution 기록을 놓치지 않게 한다.",
            "Execution은 성공과 실패 모두 DB에 반드시 기록된다.",
        ],
    )

    add_heading(doc, "9.3 UI 회귀 점검", level=2)
    add_bullets(
        doc,
        [
            "각 Phase 완료 후 브라우저에서 시나리오 3~5개를 수동 실행한다.",
            "CRUD 시나리오, 실행과 재처리 시나리오, 대시보드 기간 토글 시나리오를 포함한다.",
            "tsc --noEmit 통과는 수동 테스트의 선제 조건이다.",
        ],
    )

    add_heading(doc, "9.4 시드 기반 데모 가능성", level=2)
    add_bullets(
        doc,
        [
            "prisma/seed.ts가 8개 인터페이스와 약 60건의 실행, 약 250건의 로그를 생성한다.",
            "대시보드 차트가 의미 있게 그려지는 최소 데이터 분포를 확보한다.",
            "시드는 upsert와 deleteMany 조합으로 멱등성을 갖는다.",
        ],
    )


def section_troubleshooting(doc: Document) -> None:
    add_heading(doc, "10. 트러블슈팅 사례", level=1)

    add_heading(doc, "10.1 Prisma 7의 datasource.url 금지", level=2)
    add_paragraph(
        doc,
        "증상: npx prisma generate 실행 시 P1012 에러가 발생했다. 메시지는 'The datasource property url is no longer supported in schema files'였다.",
    )
    add_paragraph(
        doc,
        "원인: Prisma 7부터 schema의 datasource.url이 제거되고, prisma.config.ts의 datasource.url 또는 driver adapter로 넘어간다.",
    )
    add_paragraph(
        doc,
        "해결: schema.prisma에서 url 라인을 제거하고, 런타임은 lib/prisma.ts에서 PrismaLibSql을 주입하는 구조로 바꿨다. "
        "CLI용 URL은 prisma.config.ts의 datasource.url에 환경변수로 넘긴다.",
    )

    add_heading(doc, "10.2 Vercel 빌드에서 Prisma Client 파일 누락", level=2)
    add_paragraph(
        doc,
        "증상: Vercel 배포 중 'Module not found: Can''t resolve @/lib/generated/prisma/client' 에러가 발생했다.",
    )
    add_paragraph(
        doc,
        "원인: generated 파일을 .gitignore로 제외했는데, Vercel의 기본 훅만으로는 prisma generate 타이밍이 빌드 이전이 아닌 경우가 있었다.",
    )
    add_paragraph(
        doc,
        "해결: package.json에 prebuild와 postinstall 양쪽에 prisma generate를 등록했다. 어느 쪽 훅이 먼저 동작하든 Prisma Client가 준비되도록 이중화했다.",
    )

    add_heading(doc, "10.3 Turso 마이그레이션 스크립트의 SQL 파서 버그", level=2)
    add_paragraph(
        doc,
        "증상: scripts/turso-deploy.ts가 '적용 완료' 로그를 남겼는데 실제로는 'no such table' 에러가 런타임에 발생했다.",
    )
    add_paragraph(
        doc,
        "원인: Prisma가 생성한 migration.sql은 '-- CreateTable\\nCREATE TABLE ...' 형태다. 이전 구현은 세미콜론 기준으로 먼저 split한 뒤 "
        "startsWith('--') 필터를 적용했는데, 이 조합이 CREATE 문을 포함한 덩어리 전체를 통째로 제거하고 있었다.",
    )
    add_paragraph(
        doc,
        "해결: 주석 필터를 라인 단위로 선분리하도록 바꿨다. 먼저 '--'로 시작하는 라인을 제거한 뒤, 남은 SQL을 세미콜론 기준으로 분리한다. "
        "동시에 과거 실패 기록을 초기화하는 scripts/turso-reset.ts를 추가해 재시도 경로를 명확히 했다.",
    )

    add_heading(doc, "10.4 Prisma 7 모델 타입명 규약", level=2)
    add_paragraph(
        doc,
        "증상: lib/actions/*.ts에서 "
        "import type { Project, Division } from \"@/lib/generated/prisma/models\" 선언이 "
        "TS2305 'no exported member Project' 에러를 냈다.",
    )
    add_paragraph(
        doc,
        "원인: Prisma 7 Preview에서 생성되는 모델 타입에는 'Model' 접미사가 붙는다. "
        "실제 export는 ProjectModel, DivisionModel, CostItemModel 등이다.",
    )
    add_paragraph(
        doc,
        "해결: 전 파일의 import를 Model 접미사가 붙은 이름으로 교체했다. "
        "이후 같은 실수를 막기 위해 CLAUDE.md의 'Prisma 환경 메모' 섹션에 규약을 고정해 두고, "
        "다음 프로젝트에도 동일한 메모가 상속되도록 했다.",
    )

    add_heading(doc, "10.5 prisma.config.ts의 datasource.url은 CLI 경로에서 여전히 필요", level=2)
    add_paragraph(
        doc,
        "증상: prisma.config.ts에서 datasource 블록을 제거하자 npx prisma migrate dev가 "
        "'Environment variable not found: DATABASE_URL' 오류와 함께 실패했다.",
    )
    add_paragraph(
        doc,
        "원인: Prisma 7에서 deprecated된 것은 schema.prisma의 datasource.url이지, "
        "prisma.config.ts의 datasource가 아니다. CLI는 여전히 config의 datasource를 읽어 "
        "마이그레이션 대상 DB를 결정한다.",
    )
    add_paragraph(
        doc,
        "해결: schema.prisma에서는 datasource.url을 제거하고(런타임은 driver adapter가 담당), "
        "prisma.config.ts에는 datasource.url을 유지했다(CLI 전용). "
        "두 레이어가 서로 다른 역할을 맡는다는 점을 인지하는 것이 핵심이었다.",
    )
    add_paragraph(
        doc,
        "교훈: deprecated의 범위를 레이어 단위로 정확히 구분해야 한다. 릴리스 노트만 보고 "
        "섣불리 제거하면 CLI 경로가 끊긴다. 이 교훈은 CLAUDE.md에도 함께 남겼다.",
    )


def section_appendix(doc: Document) -> None:
    add_heading(doc, "11. 부록", level=1)

    add_heading(doc, "11.1 저장소와 데모", level=2)
    add_bullets(
        doc,
        [
            "저장소: https://github.com/dbwls99706/interface-hub",
            "라이브 데모: https://interface-hub.vercel.app",
        ],
    )

    add_heading(doc, "11.2 주요 디렉터리 트리", level=2)
    add_code_block(
        doc,
        "app/\n"
        "  interfaces/                인터페이스 CRUD 페이지\n"
        "    new/, [id]/, [id]/edit/\n"
        "  executions/                실행 이력과 상세 페이지\n"
        "    [id]/\n"
        "  dashboard/                 분석 대시보드 페이지\n"
        "  layout.tsx, page.tsx\n"
        "components/\n"
        "  ui/                        shadcn/ui 기반 공통 컴포넌트\n"
        "  interfaces/                인터페이스 도메인 컴포넌트\n"
        "  executions/                실행 이력 도메인 컴포넌트\n"
        "  dashboard/                 대시보드 차트/카드 컴포넌트\n"
        "lib/\n"
        "  adapters/                  프로토콜 Adapter 레이어\n"
        "  actions/                   Server Actions (CRUD, 실행 엔진, 쿼리)\n"
        "  schemas/                   Zod 스키마\n"
        "  types/                     공용 타입과 JSON 헬퍼\n"
        "  generated/prisma/          Prisma Client (빌드 시 생성)\n"
        "  prisma.ts                  PrismaClient 싱글톤\n"
        "prisma/\n"
        "  schema.prisma              데이터 모델\n"
        "  migrations/                마이그레이션 SQL\n"
        "  seed.ts                    시드 스크립트\n"
        "scripts/\n"
        "  turso-deploy.ts            Turso 마이그레이션 적용\n"
        "  turso-reset.ts             _applied_migrations 초기화\n"
        "docs/\n"
        "  screenshots/               README와 기획서용 스크린샷\n"
        "  기획서.docx, 개발문서.docx   제출용 문서",
    )

    add_heading(doc, "11.3 package.json scripts", level=2)
    add_code_block(
        doc,
        "\"scripts\": {\n"
        "  \"dev\": \"next dev\",\n"
        "  \"prebuild\": \"prisma generate\",\n"
        "  \"build\": \"next build\",\n"
        "  \"start\": \"next start\",\n"
        "  \"lint\": \"eslint\",\n"
        "  \"postinstall\": \"prisma generate\",\n"
        "  \"db:deploy\": \"tsx scripts/turso-deploy.ts\",\n"
        "  \"db:reset:remote\": \"tsx scripts/turso-reset.ts\",\n"
        "  \"db:seed:remote\": \"tsx prisma/seed.ts\"\n"
        "}",
    )

    add_heading(doc, "11.4 주요 환경변수", level=2)
    add_table(
        doc,
        headers=["변수", "용도", "예시"],
        rows=[
            ["DATABASE_URL", "로컬 SQLite 또는 Turso 단일 URL", "file:./dev.db"],
            ["TURSO_DATABASE_URL", "프로덕션(Turso) URL", "libsql://project.turso.io"],
            ["TURSO_AUTH_TOKEN", "Turso 인증 토큰", "eyJhbGciOi..."],
        ],
        col_widths_cm=[4.2, 5.6, 6.0],
    )

    add_heading(doc, "11.5 자매 프로젝트: Cost Compass", level=2)
    add_paragraph(
        doc,
        "Interface Hub의 CLAUDE.md를 상속받아 동일한 기술 스택으로 '분석 도구'를 만든 자매 사례다. "
        "환경 메모가 전파된 덕분에 Phase 0부터 Phase 3까지를 반나절에 완성했다.",
    )
    add_bullets(
        doc,
        [
            "URL: https://cost-compass.vercel.app",
            "저장소: https://github.com/dbwls99706/cost-compass",
            "의의: Interface Hub에서 겪은 환경 이슈 9건이 재발하지 않음을 확인한 실증.",
        ],
    )


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = KOR_FONT
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), KOR_FONT)
    r_fonts.set(qn("w:ascii"), KOR_FONT)
    r_fonts.set(qn("w:hAnsi"), KOR_FONT)

    build_cover(doc)
    build_toc(doc)
    section_scope(doc)
    section_architecture(doc)
    section_data_model(doc)
    section_adapter(doc)
    section_engine(doc)
    section_frontend(doc)
    section_turso(doc)
    section_vibe(doc)
    section_validation(doc)
    section_troubleshooting(doc)
    section_appendix(doc)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
