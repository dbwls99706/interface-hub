"""docs/기획서.docx 생성 스크립트."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parent
SCREENSHOTS = DOCS_DIR / "screenshots"
OUTPUT = DOCS_DIR / "기획서.docx"

KOR_FONT = "맑은 고딕"


def set_korean_font(run, size_pt: float, bold: bool = False, color=None) -> None:
    run.font.name = KOR_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), KOR_FONT)
    r_fonts.set(qn("w:ascii"), KOR_FONT)
    r_fonts.set(qn("w:hAnsi"), KOR_FONT)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: float = 11,
    bold: bool = False,
    align=None,
    space_after: float = 6,
    color=None,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_korean_font(run, size, bold=bold, color=color)
    return p


def add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    if level == 0:
        size = 22
    elif level == 1:
        size = 16
    elif level == 2:
        size = 13
    else:
        size = 12
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_korean_font(run, size, bold=True, color=RGBColor(0x1F, 0x2D, 0x3D))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.35
        run = p.runs[0] if p.runs else p.add_run("")
        # 기본 list bullet에 텍스트 채우기
        if not p.runs:
            run = p.add_run(item)
        else:
            run.text = item
        set_korean_font(run, 11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_korean_font(run, 9.5, color=RGBColor(0x55, 0x65, 0x77))
    run.italic = True


def add_image(doc: Document, path: Path, width_cm: float = 15.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    col_widths_cm: list[float] | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    if col_widths_cm:
        for col_idx, width in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[col_idx].width = Cm(width)

    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_korean_font(run, 10.5, bold=True)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            set_korean_font(run, 10.5)

    # 표 뒤 여백
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def build_cover(doc: Document) -> None:
    # 표지
    for _ in range(4):
        doc.add_paragraph()

    add_paragraph(
        doc,
        "Interface Hub",
        size=28,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x10, 0x2A, 0x43),
        space_after=4,
    )
    add_paragraph(
        doc,
        "금융사 인터페이스 통합 관리 플랫폼 기획서",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=24,
    )
    add_paragraph(
        doc,
        "노아에이티에스 2025년 상반기 연구소 인력 채용 포트폴리오",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x44, 0x55, 0x66),
        space_after=80,
    )

    add_paragraph(
        doc,
        "작성자: 홍유진 (GitHub: dbwls99706)",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "작성일: 2026-04-20",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "라이브 데모: https://interface-hub.vercel.app",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "저장소: https://github.com/dbwls99706/interface-hub",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.add_page_break()


def build_toc(doc: Document) -> None:
    add_heading(doc, "목차", level=0)
    items = [
        "1. 프로젝트 개요",
        "2. 문제 정의",
        "3. 타깃 사용자",
        "4. MVP 범위와 경계",
        "5. 핵심 기능 정의",
        "6. 성공 지표",
        "7. 확장 로드맵",
        "8. 부록",
    ]
    for item in items:
        add_paragraph(doc, item, size=12, space_after=4)
    doc.add_page_break()


def section_overview(doc: Document) -> None:
    add_heading(doc, "1. 프로젝트 개요", level=1)
    add_paragraph(
        doc,
        "Interface Hub는 금융사의 내부 핵심 시스템과 외부 기관 사이에 흩어진 다수 프로토콜 인터페이스를 "
        "단일 화면에서 관리하는 중앙화 플랫폼이다. 등록과 실행, 모니터링, 재처리, 로그 추적을 하나의 "
        "운영 콘솔로 통합해 운영팀의 인지 부하를 낮추고 장애 대응 속도를 끌어올린다.",
    )
    add_heading(doc, "1.1 프로젝트 메타", level=2)
    add_table(
        doc,
        headers=["항목", "내용"],
        rows=[
            ["프로젝트명", "Interface Hub"],
            ["라이브 데모", "https://interface-hub.vercel.app"],
            ["저장소", "https://github.com/dbwls99706/interface-hub"],
            ["개발 기간", "1일 (Claude Code 기반 바이브코딩)"],
            ["기술 스택 요약", "Next.js 16, Prisma 7, Turso libSQL, Tailwind, Recharts, Vercel"],
        ],
        col_widths_cm=[3.8, 12.0],
    )


def section_problem(doc: Document) -> None:
    add_heading(doc, "2. 문제 정의", level=1)

    add_heading(doc, "2.1 배경", level=2)
    add_paragraph(
        doc,
        "보험사, 자산운용사, 은행 등 금융사는 내부 코어 시스템과 외부 기관(금감원, 제휴사, 결제사 등) "
        "사이에 수십 개의 인터페이스를 운영한다. 이들 인터페이스의 프로토콜은 제각각이다. "
        "REST, SOAP, MQ, Batch, SFTP/FTP가 한 조직 안에서 동시에 사용된다.",
    )

    add_heading(doc, "2.2 현재의 문제점", level=2)
    add_bullets(
        doc,
        [
            "프로토콜마다 관리 도구가 흩어져 있어 운영자가 여러 콘솔을 오간다.",
            "실행 결과와 로그가 시스템별로 분산돼 있어 장애 발생 시 근본 원인 추적이 어렵다.",
            "인터페이스 추가 시마다 모니터링과 재처리 로직을 반복 구현한다.",
            "관리자 관점의 전사 운영 현황(성공률, 실패 Top 등)을 한눈에 보기 어렵다.",
        ],
    )

    add_heading(doc, "2.3 해결 방향", level=2)
    add_bullets(
        doc,
        [
            "모든 인터페이스를 단일 데이터 모델로 추상화한다 (Interface, Execution, ExecutionLog 3개 엔티티).",
            "프로토콜 차이는 Adapter 패턴으로 흡수하고, 실행 엔진은 프로토콜에 무관하게 동작한다.",
            "통합 모니터링 대시보드로 관리자 관점의 시야를 한 화면에 제공한다.",
        ],
    )


def section_users(doc: Document) -> None:
    add_heading(doc, "3. 타깃 사용자", level=1)
    add_paragraph(
        doc,
        "Interface Hub는 운영, 기획, 개발 세 축의 사용자를 동시에 고려해 설계했다. "
        "각 역할의 일상 업무에서 발생하는 반복과 마찰을 줄이는 것을 우선순위로 둔다.",
    )
    add_table(
        doc,
        headers=["역할", "주요 페인포인트", "Interface Hub가 제공하는 가치"],
        rows=[
            [
                "IT 운영팀 (1차)",
                "여러 콘솔을 오가며 실행 상태와 실패를 추적해야 한다.",
                "단일 화면에서 실행/모니터링/재처리를 모두 수행한다.",
            ],
            [
                "IT 기획/관리자 (2차)",
                "전사 인터페이스 운영 현황을 일별/주별로 정리하기 어렵다.",
                "대시보드에서 KPI와 실패율, 시계열 추이를 즉시 확인한다.",
            ],
            [
                "개발자 (3차)",
                "신규 인터페이스 추가 시 실행 로직과 로그 구조를 매번 직접 만든다.",
                "Adapter 인터페이스만 구현하면 등록과 모니터링을 자동으로 얻는다.",
            ],
        ],
        col_widths_cm=[3.5, 6.0, 6.0],
    )


def section_mvp(doc: Document) -> None:
    add_heading(doc, "4. MVP 범위와 경계", level=1)
    add_paragraph(
        doc,
        "1일이라는 짧은 개발 기간 안에 핵심 가치 가설을 검증하기 위해 범위를 명확히 잘랐다. "
        "운영 콘솔로서의 사용 흐름과 확장 가능한 구조를 동시에 보여주는 것을 우선했다.",
    )

    add_heading(doc, "4.1 MoSCoW 분석", level=2)
    add_table(
        doc,
        headers=["우선순위", "범위"],
        rows=[
            ["Must", "인터페이스 CRUD, 수동 실행, 실행 이력, 로그 조회, 분석 대시보드"],
            ["Should", "1클릭 재처리, Adapter 확장 구조, 조건부 실시간 폴링"],
            ["Could", "스케줄링, 알림 (다음 Phase)"],
            ["Won't", "인증/권한, 실제 SOAP/MQ/SFTP 클라이언트, 실시간 WebSocket (이번 MVP에서 제외)"],
        ],
        col_widths_cm=[3.2, 12.6],
    )

    add_heading(doc, "4.2 왜 Mock Adapter인가", level=2)
    add_bullets(
        doc,
        [
            "REST는 실제 외부 API(JSONPlaceholder 등)를 호출해 실동작을 증명한다.",
            "나머지 4종(SOAP, MQ, BATCH, SFTP)은 Mock Adapter로 구현해 확장 가능한 구조를 증명한다.",
            "Mock 어댑터는 SSH 핸드셰이크, WSDL 조회 같은 실제 단계 로그를 모사해 시연 시 흐름을 보여준다.",
            "신규 프로토콜 추가는 어댑터 파일 1개와 registry 1줄로 끝나도록 설계했다.",
        ],
    )


def section_features(doc: Document) -> None:
    add_heading(doc, "5. 핵심 기능 정의", level=1)
    add_paragraph(
        doc,
        "MVP가 제공하는 기능은 크게 5가지다. 각 기능은 사용자 시나리오에 맞춘 화면과 데이터 모델을 함께 갖춘다.",
    )

    add_heading(doc, "5.1 인터페이스 관리", level=2)
    add_paragraph(
        doc,
        "인터페이스를 등록하고 수정하고 삭제한다. 프로토콜별로 다른 설정(JSON config)을 동일한 폼에서 입력하며, "
        "Zod 스키마로 입력값을 컴파일/런타임 두 단계에서 검증한다. 활성/비활성 토글로 호출을 즉시 차단할 수 있다.",
    )
    add_image(doc, SCREENSHOTS / "01.png")
    add_caption(doc, "인터페이스 목록 화면. 프로토콜 뱃지와 활성 상태를 한 줄에 노출한다.")

    add_heading(doc, "5.2 실행 이력과 실시간 폴링", level=2)
    add_paragraph(
        doc,
        "모든 실행은 Execution 엔티티로 기록되며 인터페이스, 상태, 프로토콜로 필터링한다. "
        "RUNNING이나 PENDING 실행이 화면에 있을 때만 SWR이 3초 폴링을 활성화하고, 그 외에는 폴링을 멈춰 불필요한 요청을 줄인다. "
        "더 보기 버튼으로 커서 기반 페이지네이션을 제공한다.",
    )
    add_image(doc, SCREENSHOTS / "02.png")
    add_caption(doc, "실행 이력 목록. 진행 중 실행은 자동으로 새로고침된다.")

    add_heading(doc, "5.3 실행 상세와 로그 타임라인", level=2)
    add_paragraph(
        doc,
        "실행 상세는 요청, 응답, 로그를 탭으로 분리한다. 로그는 시간 순서의 타임라인 UI로 정리하고, "
        "각 로그의 metadata는 펼쳐 볼 수 있게 했다. 디버깅 시 필요한 맥락을 한 화면에서 모두 확보한다.",
    )
    add_image(doc, SCREENSHOTS / "03.png")
    add_caption(doc, "실행 상세 화면. 요청/응답/로그를 단일 페이지에서 비교한다.")

    add_heading(doc, "5.4 분석 대시보드", level=2)
    add_paragraph(
        doc,
        "최근 24시간과 7일 두 가지 기간 토글을 제공한다. KPI 카드, 시계열 BarChart, 프로토콜별 브레이크다운, "
        "실패율 Top 5, 최근 실패 10건을 한 페이지에 집계한다. 관리자 관점의 빠른 상태 점검을 지원한다.",
    )
    add_image(doc, SCREENSHOTS / "04.png")
    add_caption(doc, "분석 대시보드. 운영 현황을 단일 화면에서 요약한다.")

    add_heading(doc, "5.5 재처리 (Retry Chain)", level=2)
    add_paragraph(
        doc,
        "실패한 실행을 1클릭으로 재처리한다. 원본 Execution은 그대로 보존하고 새 Execution을 생성하며 retryOf 관계로 연결한다. "
        "재처리 체인을 따라가면 동일 인터페이스의 실패 회복 과정을 추적할 수 있다.",
    )


def section_metrics(doc: Document) -> None:
    add_heading(doc, "6. 성공 지표", level=1)

    add_heading(doc, "6.1 MVP 검증 기준", level=2)
    add_bullets(
        doc,
        [
            "8개 시드 인터페이스와 약 60건의 실행 이력이 대시보드에 정확히 집계된다.",
            "RUNNING 상태의 실행이 3초 이내 화면에 반영된다.",
            "신규 프로토콜 추가가 어댑터 파일 1개와 registry 1줄로 끝난다.",
            "로컬 SQLite와 프로덕션 Turso 사이를 환경변수만 바꾸어 전환할 수 있다.",
            "TypeScript strict 기준 npx tsc --noEmit 0건, 주요 화면 회귀 점검을 통과한다.",
        ],
    )

    add_heading(doc, "6.2 운영 KPI (제품이 운영에 투입된 가정)", level=2)
    add_table(
        doc,
        headers=["KPI", "정의", "기대 효과"],
        rows=[
            ["일 실행량", "전 인터페이스의 일별 실행 건수 합계", "트래픽 패턴 파악과 용량 산정"],
            ["프로토콜별 실패율", "(FAILED / 완료된 실행) 비율", "프로토콜 단위 안정성 모니터링"],
            ["평균 소요시간", "완료된 실행의 durationMs 평균", "성능 회귀 조기 발견"],
            ["재처리 비율", "전체 대비 retryOf가 있는 실행 비율", "운영 부담과 자동화 효과 측정"],
        ],
        col_widths_cm=[3.2, 6.0, 6.6],
    )


def section_roadmap(doc: Document) -> None:
    add_heading(doc, "7. 확장 로드맵", level=1)
    add_paragraph(
        doc,
        "MVP는 1일 기준 산출물이며 실제 운영 도구로 확장할 경로를 다음과 같이 설계한다. "
        "각 Phase는 이전 Phase의 데이터 모델과 어댑터 인터페이스를 그대로 활용한다.",
    )
    add_table(
        doc,
        headers=["Phase", "범위", "예상 공수", "선행 의존"],
        rows=[
            ["Phase 2 (M+1)", "인증/권한 (RLS, 조직 단위 분리)", "약 1주", "없음"],
            ["Phase 3 (M+2)", "실제 SOAP/MQ/SFTP 클라이언트 어댑터", "약 2주", "Phase 2"],
            ["Phase 4 (M+3)", "스케줄링 (cron 트리거)", "약 1주", "Phase 3"],
            ["Phase 5 (M+4)", "알림 연동 (Slack, 이메일)", "약 3일", "Phase 4"],
            ["Phase 6 (M+5)", "보존 정책 기반 로그 외부 저장소(S3) 연동", "약 1주", "Phase 4"],
        ],
        col_widths_cm=[3.0, 6.5, 2.6, 3.6],
    )


def section_appendix(doc: Document) -> None:
    add_heading(doc, "8. 부록", level=1)

    add_heading(doc, "8.1 용어집", level=2)
    add_table(
        doc,
        headers=["용어", "설명"],
        rows=[
            ["REST", "HTTP 기반의 자원 지향 API. JSON 요청/응답이 일반적이다."],
            ["SOAP", "XML 기반의 RPC 프로토콜. 레거시 금융 시스템에서 여전히 광범위하게 사용된다."],
            ["MQ", "메시지 큐. 비동기 메시지 발행과 구독을 제공하며 RabbitMQ, IBM MQ 등이 있다."],
            ["Batch", "주기적 대용량 처리 잡. 정산, 마감 처리에 자주 사용된다."],
            ["SFTP", "SSH 위에서 동작하는 파일 전송 프로토콜. 외부기관 파일 송수신에 자주 사용된다."],
            ["Adapter 패턴", "인터페이스를 통일하고 구현체를 갈아 끼울 수 있게 하는 설계 패턴."],
            ["SWR", "React 데이터 페칭 라이브러리. 캐시 우선 + 백그라운드 재검증 전략을 제공한다."],
        ],
        col_widths_cm=[3.2, 12.6],
    )

    add_heading(doc, "8.2 참고 자료", level=2)
    add_bullets(
        doc,
        [
            "노아에이티에스 2025년 상반기 연구소 인력 채용 공고",
            "Prisma 공식 문서: libSQL Driver Adapter",
            "Turso 공식 문서: Edge SQLite 서비스",
            "Next.js 16 공식 문서: App Router, Server Actions",
        ],
    )


def main() -> None:
    doc = Document()

    # 기본 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Normal 스타일 글꼴
    normal = doc.styles["Normal"]
    normal.font.name = KOR_FONT
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), KOR_FONT)
    r_fonts.set(qn("w:ascii"), KOR_FONT)
    r_fonts.set(qn("w:hAnsi"), KOR_FONT)

    build_cover(doc)
    build_toc(doc)
    section_overview(doc)
    section_problem(doc)
    section_users(doc)
    section_mvp(doc)
    section_features(doc)
    section_metrics(doc)
    section_roadmap(doc)
    section_appendix(doc)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
